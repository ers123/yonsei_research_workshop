"""산출물 내보내기 모듈 — 단일 '연구 설계' 결과물을 다양한 포맷으로 변환.

포함:
- create_docx : .docx (기존 app.py에 있던 함수; 이탤릭 버그 수정 후 이관)
- export_advisor_sheet : 지도교수 회의 1-pager (.md)
- export_search_queries : RISS/NTIS/Scholar 쿼리 목록 (.txt)
- export_irb_card : IRB 체크카드 (.md)
- export_bibtex_seed : BibTeX skeleton (.bib)

설계 원칙:
- 모든 변환기는 Ollama 프롬프트 한 번 호출 = 로컬 $0.
- 외부 API·DB·벡터 스토어 없음.
"""

from __future__ import annotations

import io
import re
import datetime
from typing import Callable, List

import ollama
from docx import Document
from docx.shared import Pt


# ═════════════════════════════════════════════════════════════════════
# .docx — with proper markdown rendering
# ═════════════════════════════════════════════════════════════════════

_INLINE_RE = re.compile(
    r"(\*\*[^*\n]+\*\*|__[^_\n]+__|`[^`\n]+`|\*[^*\n]+\*|_[^_\n]+_)"
)


def _add_runs_with_inline(paragraph, text: str) -> None:
    """Add runs to a paragraph, interpreting **bold**, *italic*, `code`."""
    if not text:
        return
    parts = _INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("__") and part.endswith("__"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("_") and part.endswith("_") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _split_table_cells(row_line: str) -> List[str]:
    """Parse `| a | b | c |` → ['a', 'b', 'c']."""
    s = row_line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [cell.strip() for cell in s.split("|")]


def _is_table_sep(line: str) -> bool:
    """`|---|---|` separator row."""
    s = line.strip().strip("|").replace(" ", "")
    return bool(s) and set(s) <= {"-", ":", "|"}


def _render_table(doc: Document, header: List[str], body_rows: List[List[str]]) -> None:
    """Add a Word table with bold header row."""
    cols = max(len(header), max((len(r) for r in body_rows), default=0))
    if cols == 0:
        return
    table = doc.add_table(rows=1 + len(body_rows), cols=cols)
    table.style = "Light Grid Accent 1"
    # header
    for i, cell_text in enumerate(header[:cols]):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.bold = True
    # body
    for r_idx, row in enumerate(body_rows, start=1):
        for c_idx in range(cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            text = row[c_idx] if c_idx < len(row) else ""
            _add_runs_with_inline(p, text)


def _render_markdown_to_docx(doc: Document, content: str) -> None:
    """Parse markdown-ish content and emit proper Word structures.

    Handles: # ## ### 헤딩, **bold**, *italic*, `code`, |tables|,
    - bullets, 1. numbered, > blockquote, ━/─ divider, 빈 줄.
    """
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Divider (─ or ━ blocks from director outputs)
        if stripped and set(stripped) <= {"━", "─"}:
            p = doc.add_paragraph("─" * 40)
            p.runs[0].font.color.rgb = None  # default
            i += 1
            continue

        # Empty line → paragraph break (skip)
        if not stripped:
            i += 1
            continue

        # ATX headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=4)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=2)
            i += 1
            continue

        # Blockquote (single line)
        if stripped.startswith(">"):
            text = stripped[1:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run = p.add_run(text)
            run.italic = True
            i += 1
            continue

        # Table block: "| x | y |" followed by separator
        if (stripped.startswith("|") and stripped.endswith("|")
                and i + 1 < len(lines) and _is_table_sep(lines[i + 1])):
            header = _split_table_cells(stripped)
            body: List[List[str]] = []
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith("|") and lines[j].strip().endswith("|"):
                if _is_table_sep(lines[j]):
                    j += 1
                    continue
                body.append(_split_table_cells(lines[j].strip()))
                j += 1
            _render_table(doc, header, body)
            i = j
            continue

        # Unordered list
        if re.match(r"^\s*[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            text = re.sub(r"^\s*[-*]\s+", "", line)
            _add_runs_with_inline(p, text)
            i += 1
            continue

        # Ordered list
        if re.match(r"^\s*\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            _add_runs_with_inline(p, text)
            i += 1
            continue

        # Regular paragraph — inline formatting
        p = doc.add_paragraph()
        _add_runs_with_inline(p, line)
        i += 1


def create_docx(content: str, header: str, disclaimer: str) -> io.BytesIO:
    """Render final research-design text into a .docx file with proper
    markdown parsing: headings, bold/italic/code, tables, lists, blockquotes.
    """
    doc = Document()
    doc.add_heading(header, level=1)
    doc.add_paragraph(
        f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )

    # Italic disclaimer
    p = doc.add_paragraph()
    run = p.add_run(disclaimer)
    run.italic = True
    doc.add_paragraph("")

    _render_markdown_to_docx(doc, content)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ═════════════════════════════════════════════════════════════════════
# Ollama-backed generators (all non-streaming; we want full output)
# ═════════════════════════════════════════════════════════════════════

_EMPTY_CONTENT_HINT = (
    "[빈 응답 — 사고 모드가 켜진 상태에서 예산 내에 최종 응답을 내지 못했을 수 있습니다. "
    "사고 모드를 끄거나 더 작은 산출물부터 시도하세요.]"
)


def _ollama_once(model: str, system_prompt: str, user_payload: str, max_tokens: int = 2048,
                 think: bool = False, keep_alive: int = 0) -> str:
    """Single non-streaming ollama.chat call — returns full content.

    think=False by default (exports don't need reasoning, just structured output).
    keep_alive=0 unloads the model after the call for VRAM hygiene."""
    try:
        response = ollama.chat(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_payload},
            ],
            think=think,
            options={"temperature": 0.2, "top_p": 0.9, "num_predict": max_tokens},
            keep_alive=keep_alive,
            stream=False,
        )
        msg = response.get("message") if isinstance(response, dict) else getattr(response, "message", None)
        if msg is None:
            return _EMPTY_CONTENT_HINT
        content = msg.get("content", "") if isinstance(msg, dict) else (getattr(msg, "content", "") or "")
        if not content.strip():
            return _EMPTY_CONTENT_HINT
        return content
    except Exception as e:
        return f"[변환 실패: {e}]"


# --- 지도교수 회의 1-pager ---

_ADVISOR_SHEET_PROMPT = """당신은 대학원생이 '내일 아침 지도교수 미팅'에 들고 갈 1페이지 요약서를 작성합니다.
제공된 연구 설계 전체를 짧게 압축하고, 교수가 던질 법한 질문과 답을 준비합니다.

출력은 순수 Markdown, 아래 구조를 그대로 따르세요:

# 지도교수 회의 준비 1-pager

## 1. Elevator Pitch (한 문장)
> ...

## 2. Research Question
- RQ:
- 가설 (H1, H2):

## 3. 핵심 설계 요약 (표 1개)
| 방법론 | 대상 | 데이터 | 기간 |
|---|---|---|---|
| | | | |

## 4. 예상 질문 × 답 준비 (5개)
| # | 예상 질문 | 답변 요지 (1-2문장) |
|---|-----------|----------------------|

## 5. 오늘 지도교수에게 받고 싶은 결정 1개
> ...

## 6. 남은 리스크 2개
- ...
- ...
"""


def export_advisor_sheet(model: str, director_output: str, advisor_output: str = "") -> str:
    """지도교수 회의 1페이지 생성. advisor_output이 있으면 함께 참고."""
    payload = (
        f"=== 최종 연구 설계 ===\n{director_output}\n=== 끝 ===\n\n"
        f"=== (선택) 지도교수 에이전트 출력 ===\n{advisor_output or '(없음)'}\n=== 끝 ==="
    )
    return _ollama_once(model, _ADVISOR_SHEET_PROMPT, payload, max_tokens=1500)


# --- RISS / NTIS / Scholar 검색 쿼리 ---

_SEARCH_QUERIES_PROMPT = """당신은 대학원생이 문헌 탐색에 바로 사용할 검색 쿼리를 생성합니다.
제공된 연구 설계 전체를 바탕으로 국내/국외 DB용 쿼리를 만듭니다.

출력은 순수 텍스트. 아래 섹션을 그대로:

# 문헌 검색 쿼리 모음

## RISS (국내 한국어, 국문 키워드)
1. "<쿼리 문자열>"
2. ...
(10개)

## NTIS (국가과제 DB, 한국어)
1. ...
(5개)

## Google Scholar (영문)
1. <쿼리>
2. ...
(10개)

## 불리언 조합 (고급 검색용)
1. (keyword1 OR keyword2) AND (keyword3) NOT (keyword4)
2. ...
(5개)

원칙:
- 따옴표·AND/OR/NOT을 적극 활용
- 너무 광범위/협소하지 않게
- 각 쿼리 옆에 간단한 의도 주석 달아도 됨"""


def export_search_queries(model: str, director_output: str) -> str:
    payload = f"=== 연구 설계 ===\n{director_output}\n=== 끝 ==="
    return _ollama_once(model, _SEARCH_QUERIES_PROMPT, payload, max_tokens=1500)


# --- IRB 체크카드 ---

_IRB_CARD_PROMPT = """당신은 대학원생이 IRB(연구윤리심의) 신청 전 자가점검할 체크카드를 만듭니다.
제공된 연구 설계를 바탕으로 10개 항목을 생성합니다.
각 항목은 '예/아니오/해당없음' 판단 가능하도록 구체적이어야 합니다.

출력은 순수 Markdown:

# IRB 사전 체크카드

연구 제목: (1줄)

| # | 점검 항목 | 해당 여부 | 준비 문서 |
|---|-----------|-----------|-----------|
| 1 | 인간대상 연구 여부 확인 | 예/아니오 | ... |
| 2 | ... | | |

(총 10개 — 최소한 아래 주제를 커버: 연구대상자 모집, 동의서, 개인정보 처리, 민감정보 여부, 취약집단, 위험·이익 분석, 데이터 저장·폐기, 연구자료 공유, 2차 이용, 연구윤리 교육 이수)

## 요약 판단
- 심의 수준 추정: (신속/정규/면제 중 하나 + 이유 1-2줄)
- 가장 주의해야 할 항목 TOP 3:
  1. ...
  2. ...
  3. ...
"""


def export_irb_card(model: str, director_output: str) -> str:
    payload = f"=== 연구 설계 ===\n{director_output}\n=== 끝 ==="
    return _ollama_once(model, _IRB_CARD_PROMPT, payload, max_tokens=1800)


# --- BibTeX skeleton ---

_BIBTEX_PROMPT = """당신은 대학원생의 참고문헌 관리 초기 작업을 돕는 BibTeX skeleton을 생성합니다.
제공된 연구 설계에서 키워드·주제·방법론·선행연구 후보를 뽑아
검색해보면 존재할 법한 논문의 '추정(skeleton)' BibTeX 엔트리를 10개 만듭니다.

★ 중요: 이 엔트리는 모두 실제 존재 확인이 안 된 stub입니다.
따라서 모든 엔트리에 note 필드로 "LLM-generated stub — 실제 존재 검증 필요"를 명시하세요.
존재 확률이 높은 것부터 낮은 순으로 정렬하세요.

출력은 순수 BibTeX 문법:

@misc{stub_kim_2023,
  author = {Kim, ...},
  title  = {...},
  year   = {2023},
  note   = {LLM-generated stub — 실제 존재 검증 필요},
  keywords = {...}
}

@misc{...}

(10개)

각 엔트리 위에 Bash 주석(% ...)으로 '추정 근거' 한 줄씩 추가."""


def export_bibtex_seed(model: str, director_output: str) -> str:
    payload = f"=== 연구 설계 ===\n{director_output}\n=== 끝 ==="
    return _ollama_once(model, _BIBTEX_PROMPT, payload, max_tokens=2000)


# ═════════════════════════════════════════════════════════════════════
# Export 디스패치 (UI에서 참조)
# ═════════════════════════════════════════════════════════════════════

# key → (label, filename suffix, mime, generator signature)
EXPORT_REGISTRY = {
    "docx": {
        "label": "Word 문서 (.docx)",
        "filename": "연구설계초안",
        "ext": "docx",
        "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "kind": "docx",
    },
    "advisor_sheet": {
        "label": "지도교수 회의 1-pager (.md)",
        "filename": "지도교수회의_1pager",
        "ext": "md",
        "mime": "text/markdown",
        "kind": "llm",
        "fn": export_advisor_sheet,
    },
    "search_queries": {
        "label": "RISS/NTIS/Scholar 검색 쿼리 (.txt)",
        "filename": "문헌검색쿼리",
        "ext": "txt",
        "mime": "text/plain",
        "kind": "llm",
        "fn": export_search_queries,
    },
    "irb_card": {
        "label": "IRB 체크카드 (.md)",
        "filename": "IRB_체크카드",
        "ext": "md",
        "mime": "text/markdown",
        "kind": "llm",
        "fn": export_irb_card,
    },
    "bibtex_seed": {
        "label": "BibTeX seed (.bib)",
        "filename": "bibtex_seed",
        "ext": "bib",
        "mime": "text/plain",
        "kind": "llm",
        "fn": export_bibtex_seed,
    },
}
